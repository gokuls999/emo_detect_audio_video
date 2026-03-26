"""Generate MSTDN-A Implementation Progress Report as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helper utilities ──────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p

def body(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, bold=bold)
    return p

def bullet(text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    p = doc.add_paragraph(style=style)
    if sub:
        p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=10, color=(255, 255, 255))
        # Header background colour
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F497D")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        fill = "DCE6F1" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), fill)
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def divider():
    p = doc.add_paragraph("─" * 85)
    p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
    p.runs[0].font.size = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("MSTDN-A Implementation Progress Report")
set_font(run, size=20, bold=True, color=(31, 73, 125))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run("Multimodal Speech-Temporal Distillation Network (Audio-Primary)")
set_font(run, size=13, color=(89, 89, 89))

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run(f"Researcher: Bineetha   |   Date: {date.today().strftime('%B %d, %Y')}   |   Status: Stage 2 Training In Progress")
set_font(run, size=10, color=(128, 128, 128))

doc.add_paragraph()
divider()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
body(
    "MSTDN-A is a PhD-level speech emotion and stress detection system. "
    "The core research claim is that a student audio model, trained under "
    "physiological supervision (EEG / GSR / PPG) via knowledge distillation, "
    "can estimate affective states that are physiologically grounded — from room "
    "audio alone — without wearable sensors at inference time."
)
doc.add_paragraph()
add_table(
    ["Property", "Value"],
    [
        ["Full system name",   "MSTDN-A — Multimodal Speech-Temporal Distillation Network"],
        ["Primary input",      "Room microphone (audio-centric)"],
        ["GPU",                "NVIDIA GeForce GTX 1080 Ti  (11 GB VRAM, SM 6.1)"],
        ["Framework",          "PyTorch 2.7.1+cu118  |  Python 3.11"],
        ["Dataset",            "MAFW — 47 subjects, 10,045 clips, 11 emotion classes"],
        ["Training split",     "single/no_caption / set_1  —  7,333 training clips"],
        ["Current stage",      "Stage 2 — Multimodal Teacher Training"],
    ],
    col_widths=[2.2, 4.0],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ENVIRONMENT SETUP
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Environment Setup & GPU Compatibility Fix")
body(
    "The GTX 1080 Ti uses CUDA compute capability SM 6.1 (Pascal architecture). "
    "PyTorch 2.10+cu128 requires a minimum of SM 7.0, making it incompatible with this GPU. "
    "The following solution was identified and applied:"
)
doc.add_paragraph()
add_table(
    ["Issue", "Root Cause", "Solution"],
    [
        ["PyTorch CUDA mismatch",
         "cu128 wheels require SM ≥ 7.0; GTX 1080 Ti is SM 6.1",
         "Installed PyTorch 2.7.1+cu118 via Python 3.11 (cu118 supports SM 6.1)"],
        ["torchaudio.load failing",
         "torchcodec requires FFmpeg DLLs not available on Windows",
         "Replaced all audio I/O with soundfile + librosa.resample"],
        ["HuggingFace offline failure",
         "No internet — model tried to re-download on each run",
         "Set HF_HUB_OFFLINE=1; models already cached locally"],
    ],
    col_widths=[1.6, 2.5, 2.2],
)
doc.add_paragraph()
body("Confirmed GPU execution result:", bold=True)
bullet("Device: NVIDIA GeForce GTX 1080 Ti  |  VRAM: 11.0 GB")
bullet("Forward pass: primary_logits (1×11), embedding Z_s (1×256) — all on cuda:0")
bullet("VRAM used: 403 MB  |  Status: SUCCESS")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 3. CODE BUGS FIXED
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Code Bugs Identified and Fixed")
add_table(
    ["File", "Bug", "Fix Applied"],
    [
        ["training/common.py",
         "load_config() did plain yaml.safe_load — stage configs with 'inherits: base.yaml' lost all base keys",
         "Added _deep_merge() helper; load_config() now loads parent YAML and merges overrides recursively"],
        ["training/common.py",
         "resolve_device() did not verify CUDA actually works before returning cuda device",
         "Added torch.zeros(1).cuda() probe with RuntimeError catch; falls back to CPU safely"],
        ["data/dataset.py",
         "annotation.xlsx uses -1 for 'not rated' fields; these were passed as probability targets",
         "Added np.clip(values, 0.0, 1.0) to clamp -1 values to 0 before use"],
        ["data/dataset.py  &  data/physio_sync.py",
         "torchaudio.load / torchaudio.info calls fail due to missing FFmpeg DLLs",
         "Replaced with sf.read() / sf.info() from soundfile library"],
        ["models/teacher/caption_encoder.py",
         "AutoModel loads full CLIPModel (vision+text); calling it with text-only tokens raises ValueError",
         "Switched to CLIPTextModel.from_pretrained() — text encoder only, no pixel_values needed"],
        ["models/student/prosodic_branch.py",
         "Conv1d expected in_channels=20 but build_prosodic_proxy generates 40 channels (20 chunks × 2 stats)",
         "Changed default in_channels parameter from 20 to 40"],
        ["training/stage2_teacher.py",
         "Checkpoint saved only after all 50 epochs — any interruption loses all progress",
         "Added per-epoch checkpoint save (stage2_resume.pt) with model + optimizer state + epoch number; auto-resumes on restart"],
    ],
    col_widths=[1.7, 2.4, 2.2],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. TRAINING VERIFICATION
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Training Pipeline Verification")
body("A 3-batch smoke test was run using real data to confirm the full pipeline before full training:")
doc.add_paragraph()
add_table(
    ["Batch", "Total Loss", "CE Loss", "Stress Loss", "Time", "VRAM"],
    [
        ["1", "3.1125", "2.5008", "0.1739", "0.59 s", "301 MB"],
        ["2", "3.3382", "2.5648", "0.1783", "0.05 s", "301 MB"],
        ["3", "3.1848", "2.6880", "0.1282", "0.05 s", "301 MB"],
    ],
    col_widths=[0.7, 1.1, 1.0, 1.1, 0.9, 0.9],
)
doc.add_paragraph()
body("Observations:", bold=True)
bullet("Cross-entropy ≈ 2.5 is exactly expected for a randomly initialised model over 11 classes (log(11) ≈ 2.398)")
bullet("Backward pass and optimizer step complete without error — gradients flow correctly")
bullet("301 MB VRAM usage is well within the 11 GB available")
bullet("Batch 1 is slow (model loading); subsequent batches run at ~4 batches/second")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CURRENT TRAINING STATUS
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Current Training Status — Stage 2 Teacher")
add_table(
    ["Parameter", "Value"],
    [
        ["Stage",            "Stage 2 — Multimodal Teacher Training"],
        ["Model",            "MSTDNATeacher (EEG + GSR + PPG + Audio + CLIP Text)"],
        ["Epochs",           "50 total"],
        ["Batch size",       "16"],
        ["Batches/epoch",    "459"],
        ["Speed",            "~4 batches/second"],
        ["Time/epoch",       "~2 minutes"],
        ["Total duration",   "~100 minutes (~1 hour 40 min)"],
        ["Checkpoint",       "checkpoints/stage2_resume.pt  (saved after every epoch)"],
        ["Final output",     "checkpoints/stage2_teacher.pt  (saved after epoch 50)"],
        ["Resume command",   "HF_HUB_OFFLINE=1 py -3.11 -m training.stage2_teacher --config configs/stage2_teacher.yaml"],
    ],
    col_widths=[2.0, 4.3],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. REMAINING TRAINING PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Remaining Training Pipeline")
add_table(
    ["Stage", "Name", "Description", "Input", "Output"],
    [
        ["Stage 2", "Teacher Training",
         "Train full multimodal teacher on EEG + GSR + PPG + audio + captions with multi-task loss",
         "EEG, GSR, PPG, audio, captions",
         "stage2_teacher.pt"],
        ["Stage 3", "Student Distillation",
         "Freeze teacher; train audio-only student with task loss + MSE(Z_s, Z_t) + RKD + InfoNCE",
         "Audio only",
         "stage3_student.pt"],
        ["Stage 4", "Caption Refinement",
         "CLIP text alignment on 8034 captioned clips; anchors acoustic embeddings to semantic descriptions",
         "Audio + captions",
         "stage4_student.pt"],
        ["Stage 5", "Domain Augmentation",
         "Room impulse response, noise injection, mic degradation, pitch/speed shift",
         "Augmented audio",
         "stage5_student.pt"],
    ],
    col_widths=[0.7, 1.5, 2.5, 1.3, 1.3],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 7. MODEL OUTPUT HEADS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Model Output Heads")
body("Each 3-second audio window produces the following per-speaker outputs:")
doc.add_paragraph()
add_table(
    ["Output", "Type", "Range", "Grounding"],
    [
        ["primary_emotion",      "Softmax classifier",   "11 classes",  "Annotated single label"],
        ["emotion_distribution", "Probability vector",   "[0,1]^11",    "Multi-annotator soft labels"],
        ["secondary_emotions",   "Multi-label sigmoid",  "[0,1]^11",    "Multi-label annotations"],
        ["valence",              "Regression",           "[-1, +1]",    "Derived from emotion columns"],
        ["arousal",              "Regression",           "[0, 1]",      "Derived from high-energy emotions"],
        ["stress_score",         "Regression",           "[0, 1]",      "GSR SCR-anchored during training"],
    ],
    col_widths=[1.7, 1.5, 1.1, 2.0],
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Next Steps")
bullet("Complete Stage 2 training (resume with HF_HUB_OFFLINE=1 command above)")
bullet("Run Stage 3 — student distillation training (audio only, teacher frozen)")
bullet("Run Stage 4 — caption-guided acoustic refinement")
bullet("Run Stage 5 — acoustic domain augmentation")
bullet("Evaluate on all 5-fold CV splits + Leave-One-Subject-Out (LOSO)")
bullet("Ablation study: prosodic-only vs spectral-only vs deep audio vs full MSTDN-A")
bullet("Baseline comparison: MFCCs+LSTM, wav2vec2 fine-tune, EmoFusion, openSMILE+SVM")
bullet("Paper writing — IEEE Trans. Affective Computing (primary venue)")

doc.add_paragraph()
divider()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run("MSTDN-A — PhD Research  |  Confidential  |  " + date.today().strftime("%B %Y"))
set_font(run, size=9, color=(128, 128, 128))

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "C:/Users/ADMIN/Desktop/Bineetha - emoDet/MSTDN_A_Progress_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
